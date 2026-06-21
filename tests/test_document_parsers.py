from pathlib import Path

from docx import Document

from feature_keyword_extractor.parsers.docx_parser import DocxParser
from feature_keyword_extractor.parsers.markdown_parser import MarkdownParser


def test_docx_parser_builds_heading_dependency_graph(tmp_path: Path):
    docx_path = tmp_path / "需求文档A.docx"
    doc = Document()
    doc.add_heading("账号体系", level=1)
    doc.add_paragraph("账号能力总述。")
    doc.add_heading("登录管理", level=2)
    doc.add_heading("手机号登录", level=3)
    doc.add_paragraph("用户输入手机号后，系统发送短信验证码。")
    doc.add_paragraph("验证码校验成功后写入登录态。")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "手机号"
    doc.save(docx_path)

    parsed = DocxParser().parse(docx_path)

    assert parsed.doc_id == "需求文档A.docx"
    phone = next(node for node in parsed.heading_nodes if node.title == "手机号登录")
    parent = next(node for node in parsed.heading_nodes if node.heading_id == phone.parent_id)

    assert parent.title == "登录管理"
    assert phone.ancestor_titles == ["账号体系", "登录管理"]
    assert len(phone.body_ids) == 3
    assert "level" not in phone.model_dump()

    body_text = "\n".join(block.text for block in parsed.body_blocks if block.body_id in phone.body_ids)
    assert "短信验证码" in body_text
    assert "字段\t手机号" in body_text


def test_markdown_parser_keeps_fourth_level_content_under_nearest_heading(tmp_path: Path):
    md_path = tmp_path / "需求文档B.md"
    md_path.write_text(
        "# 账号体系\n"
        "账号能力总述。\n\n"
        "## 登录管理\n"
        "### 手机号登录\n"
        "用户输入手机号后获取验证码。\n\n"
        "#### 异常处理\n"
        "验证码过期后需要重新发送。\n",
        encoding="utf-8",
    )

    parsed = MarkdownParser().parse(md_path)

    phone = next(node for node in parsed.heading_nodes if node.title == "手机号登录")
    assert phone.ancestor_titles == ["账号体系", "登录管理"]
    assert len(phone.body_ids) == 2

    body_text = "\n".join(block.text for block in parsed.body_blocks if block.body_id in phone.body_ids)
    assert "用户输入手机号" in body_text
    assert "异常处理" in body_text
    assert "验证码过期" in body_text
